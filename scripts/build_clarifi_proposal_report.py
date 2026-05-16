from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "Team7_ClariFi_Proposal_Report.docx"
PDF_OUT = ROOT / "reports" / "Team7_ClariFi_Proposal_Report.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_question(doc, label, question, answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(f"{label}. {question} ")
    style_run(r, bold=True, color="1F4D78")
    p.add_run(answer)


def add_compact_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    style_run(r, bold=True, color="2E74B5", size=12)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.05

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ClariFi: Interactive Visual Analytics for Mortgage Readiness")
style_run(r, bold=True, color="0B2545", size=16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Team 7 | Lalitha Dasu, Vandana Mansur, Pranav Manimaran | ECS 273 Visual Analytics | Spring 2026")
style_run(r, color="4B5563", size=10.5)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(5)
lead.add_run("Proposal focus. ").bold = True
lead.add_run(
    "ClariFi helps first-time homebuyers understand mortgage readiness through objective, data-driven visual analytics. "
    "The narrowed scope uses 2025 California HMDA loan records as the main large dataset, with optional user transaction CSVs and BLS Consumer Expenditure benchmarks for personal spending context. "
    "The system combines non-trivial computation, explainable ML, and linked interactive views rather than giving a simple calculator result."
)

add_compact_heading(doc, "Heilmeier Questions")
questions = [
    (
        "Q1",
        "What are you trying to do?",
        "We will show users how close they are to mortgage readiness by comparing their income, debt, loan amount, property value, and location with real California mortgage applications."
    ),
    (
        "Q2",
        "How is it done today, and what are the limits?",
        "Current calculators and finance apps usually return generic debt-to-income rules, yes/no prequalification, or chatbot advice. They rarely expose peer borrower distributions, geographic approval patterns, model uncertainty, or why a scenario changes the outcome."
    ),
    (
        "Q3",
        "What is new, and why will it succeed?",
        "ClariFi links an XGBoost approval-likelihood model, SHAP explanations, counterfactual what-if controls, and coordinated dashboards. It should succeed because each recommendation is tied to large public loan records and visible evidence rather than unsupported advice."
    ),
    (
        "Q4",
        "Who cares?",
        "Primary users are renters and first-time buyers comparing affordability across counties. Secondary users include housing counselors, fintech product teams, and instructors evaluating explainable visual analytics on real-world data."
    ),
    (
        "Q5",
        "What impact will it make, and how will you measure it?",
        "Impact is measured by model quality (AUC, F1, balanced accuracy, calibration), explanation usefulness (whether users can identify top approval drivers), and visualization tasks such as finding a county, comparing borrower peers, and testing one feasible change."
    ),
    (
        "Q6",
        "What are the risks and payoffs?",
        "Risks include HMDA missingness, bias in historical lending decisions, and users over-trusting predictions. We mitigate these with filters, disclaimers, calibration, and SHAP/uncertainty displays. Payoff is a transparent mortgage-readiness tool and strong course portfolio artifact."
    ),
    (
        "Q7",
        "How much will it cost?",
        "The project cost is zero dollars: HMDA, BLS, and Census data are public; modeling uses Python, XGBoost, SHAP, and scikit-learn; the dashboard uses open-source React/FastAPI tools."
    ),
    (
        "Q8",
        "How long will it take?",
        "Five weeks, from Apr. 28 to Jun. 6, 2026."
    ),
    (
        "Q9",
        "What are the midterm and final exams?",
        "Midterm: clean HMDA sample, baseline model, EDA, and two linked prototype views. Final: tuned model with SHAP, dashboard with map/scatter/what-if controls, exported model outputs, short user evaluation, report, and demo."
    ),
]
for q in questions:
    add_question(doc, *q)

add_compact_heading(doc, "Datasets and Computation")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run("Main dataset: ").bold = True
p.add_run(
    "2025 HMDA Modified Loan/Application Register, published by CFPB/FFIEC, with loan-level records from about 4,768 filers. "
    "For course feasibility we use a 60,000-row California model-ready sample filtered to home-purchase, owner-occupied, first-lien applications. "
    "Fields include action taken, county/MSA, income, loan amount, property value, debt-to-income ratio, combined loan-to-value ratio, loan type, applicant age, sex, race, ethnicity, and census-tract income measures. "
)
p.add_run("Computation: ").bold = True
p.add_run(
    "data cleaning, missingness analysis, county aggregation, nonlinear feature ranking using mutual information, XGBoost model selection, threshold tuning, calibration, SHAP global/local explanations, and feasible counterfactual scenario changes."
)

add_compact_heading(doc, "Expected Innovations")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run("Algorithm/computation: ").bold = True
p.add_run(
    "ClariFi combines approval likelihood, SHAP feature attribution, DTI/LTV risk surfaces, and what-if scenario search over controllable features such as down payment, debt ratio, and target county. "
)
p.add_run("Visualization: ").bold = True
p.add_run(
    "It uses linked county choropleths, borrower scatterplots, benchmark bars, histograms, and AI-guided highlights so model outputs become explorable visual evidence."
)

add_compact_heading(doc, "Brief Literature Survey")
lit = [
    ["Fuster et al. (2022)", "ML changes credit-market predictions and can affect groups unevenly.", "Supports mortgage ML framing; limitation: policy analysis, not user-facing visual guidance."],
    ["Yuan et al. (2021)", "Reviews visual analytics techniques for ML: data quality, feature selection, and model understanding.", "Guides our EDA/model-debug views; limitation: general survey, not finance-specific."],
    ["Dimara & Perin (2020)", "Defines interaction for visualization and motivates richer filtering/what-if design.", "Supports brushing, linking, and scenario controls; limitation: conceptual, not domain implementation."],
    ["Alicioglu & Sun (2021)", "Surveys visual analytics approaches for explainable AI.", "Guides SHAP explanation views; limitation: broad XAI survey, not mortgage readiness."],
    ["Liao & Varshney (2021)", "Argues XAI must be designed around user goals and explanations, not only algorithms.", "Supports plain-language AI agent explanations; limitation: does not prescribe a full dashboard architecture."],
    ["Mahajan et al. (2019)", "Adds feasibility constraints to counterfactual explanations.", "Informs realistic what-if changes; limitation: counterfactual generation can be complex, so we restrict to controllable finance features."],
]
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)
hdr = table.rows[0].cells
for i, text in enumerate(["Paper", "Main idea", "Use and limitation for ClariFi"]):
    hdr[i].text = text
    set_cell_shading(hdr[i], "F4F6F9")
    set_cell_margins(hdr[i])
    for paragraph in hdr[i].paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(8.5)
for row in lit:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(8.2)
table.columns[0].width = Inches(1.25)
table.columns[1].width = Inches(2.25)
table.columns[2].width = Inches(3.0)

add_compact_heading(doc, "Plan of Activities")
plan = [
    ["Activity", "Owner(s)", "Dates"],
    ["Scope, data inventory, dashboard sketch", "All", "Apr 28-May 4"],
    ["HMDA/BLS/transaction pipeline and EDA", "Lalitha, Vandana", "May 5-11"],
    ["XGBoost, SHAP, risk surfaces, what-if logic", "Lalitha, Vandana", "May 12-18"],
    ["Linked dashboard: map, scatter, benchmark views", "Pranav, Vandana", "May 19-25"],
    ["Agent explanations, evaluation, final report/demo", "All", "May 26-Jun 6"],
]
plan_table = doc.add_table(rows=1, cols=3)
plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(plan_table)
for i, text in enumerate(plan[0]):
    plan_table.rows[0].cells[i].text = text
    set_cell_shading(plan_table.rows[0].cells[i], "F4F6F9")
for row in plan[1:]:
    cells = plan_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
for row in plan_table.rows:
    for cell in row.cells:
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(8.8)
                if row is plan_table.rows[0]:
                    run.bold = True

doc.add_section(WD_SECTION.NEW_PAGE)
add_compact_heading(doc, "References")
refs = [
    "Alicioglu, G., & Sun, B. (2021). A survey of visual analytics for Explainable Artificial Intelligence methods. Computers & Graphics, 102, 502-520. https://doi.org/10.1016/j.cag.2021.09.002",
    "Consumer Financial Protection Bureau. (2026). 2025 HMDA Data on Mortgage Lending Now Available. https://www.consumerfinance.gov/about-us/newsroom/2025-hmda-data-on-mortgage-lending-now-available/",
    "Dimara, E., & Perin, C. (2020). What is Interaction for Data Visualization? IEEE Transactions on Visualization and Computer Graphics, 26(1), 119-129. https://doi.org/10.1109/TVCG.2019.2934283",
    "Fuster, A., Goldsmith-Pinkham, P., Ramadorai, T., & Walther, A. (2022). Predictably Unequal? The Effects of Machine Learning on Credit Markets. Journal of Finance, 77(1), 5-47. https://doi.org/10.1111/jofi.13090",
    "Liao, Q. V., & Varshney, K. R. (2021). Human-Centered Explainable AI (XAI): From Algorithms to User Experiences. arXiv:2110.10790. https://arxiv.org/abs/2110.10790",
    "Mahajan, D., Tan, C., & Sharma, A. (2019). Preserving Causal Constraints in Counterfactual Explanations for Machine Learning Classifiers. NeurIPS CausalML Workshop. https://www.microsoft.com/en-us/research/publication/preserving-causal-constraints-in-counterfactual-explanations-for-machine-learning-classifiers/",
    "Yuan, J., Chen, C., Yang, W., Liu, M., Xia, J., & Liu, S. (2021). A survey of visual analytics techniques for machine learning. Computational Visual Media, 7, 3-36. https://doi.org/10.1007/s41095-020-0191-7",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    for run in p.runs:
        run.font.size = Pt(10.5)

doc.core_properties.title = "ClariFi Proposal Report"
doc.core_properties.author = "Team 7"
doc.core_properties.subject = "ECS 273 Visual Analytics Final Project Proposal"
doc.save(OUT)
print(OUT)


pdf_styles = getSampleStyleSheet()
pdf_styles.add(ParagraphStyle(
    name="ReportTitle",
    parent=pdf_styles["Title"],
    fontName="Helvetica-Bold",
    fontSize=15.5,
    leading=18,
    textColor=colors.HexColor("#0B2545"),
    alignment=TA_CENTER,
    spaceAfter=3,
))
pdf_styles.add(ParagraphStyle(
    name="Meta",
    parent=pdf_styles["Normal"],
    fontName="Helvetica",
    fontSize=9,
    leading=11,
    textColor=colors.HexColor("#4B5563"),
    alignment=TA_CENTER,
    spaceAfter=5,
))
pdf_styles.add(ParagraphStyle(
    name="Compact",
    parent=pdf_styles["Normal"],
    fontName="Helvetica",
    fontSize=8.8,
    leading=10.5,
    spaceAfter=3,
))
pdf_styles.add(ParagraphStyle(
    name="Heading",
    parent=pdf_styles["Heading2"],
    fontName="Helvetica-Bold",
    fontSize=11.2,
    leading=12.5,
    textColor=colors.HexColor("#2E74B5"),
    spaceBefore=5,
    spaceAfter=3,
))
pdf_styles.add(ParagraphStyle(
    name="Ref",
    parent=pdf_styles["Normal"],
    fontName="Helvetica",
    fontSize=9.6,
    leading=11,
    leftIndent=18,
    firstLineIndent=-18,
    spaceAfter=5,
))
pdf_styles.add(ParagraphStyle(
    name="SmallTable",
    parent=pdf_styles["Normal"],
    fontName="Helvetica",
    fontSize=7.6,
    leading=8.6,
))


story = [
    Paragraph("ClariFi: Interactive Visual Analytics for Mortgage Readiness", pdf_styles["ReportTitle"]),
    Paragraph("Team 7 | Lalitha Dasu, Vandana Mansur, Pranav Manimaran | ECS 273 Visual Analytics | Spring 2026", pdf_styles["Meta"]),
    Paragraph(
        "<b>Proposal focus.</b> ClariFi helps first-time homebuyers understand mortgage readiness through objective, data-driven visual analytics. "
        "The narrowed scope uses 2025 California HMDA loan records as the main large dataset, with optional user transaction CSVs and BLS Consumer Expenditure benchmarks for personal spending context. "
        "The system combines non-trivial computation, explainable ML, and linked interactive views rather than giving a simple calculator result.",
        pdf_styles["Compact"],
    ),
    Paragraph("Heilmeier Questions", pdf_styles["Heading"]),
]
for label, question, answer in questions:
    story.append(Paragraph(f"<b><font color='#1F4D78'>{label}. {question}</font></b> {answer}", pdf_styles["Compact"]))

story.extend([
    Paragraph("Datasets and Computation", pdf_styles["Heading"]),
    Paragraph(
        "<b>Main dataset:</b> 2025 HMDA Modified Loan/Application Register, published by CFPB/FFIEC, with loan-level records from about 4,768 filers. "
        "For course feasibility we use a 60,000-row California model-ready sample filtered to home-purchase, owner-occupied, first-lien applications. "
        "Fields include action taken, county/MSA, income, loan amount, property value, debt-to-income ratio, combined loan-to-value ratio, loan type, applicant age, sex, race, ethnicity, and census-tract income measures. "
        "<b>Computation:</b> data cleaning, missingness analysis, county aggregation, nonlinear feature ranking using mutual information, XGBoost model selection, threshold tuning, calibration, SHAP global/local explanations, and feasible counterfactual scenario changes.",
        pdf_styles["Compact"],
    ),
    Paragraph("Expected Innovations", pdf_styles["Heading"]),
    Paragraph(
        "<b>Algorithm/computation:</b> ClariFi combines approval likelihood, SHAP feature attribution, DTI/LTV risk surfaces, and what-if scenario search over controllable features such as down payment, debt ratio, and target county. "
        "<b>Visualization:</b> It uses linked county choropleths, borrower scatterplots, benchmark bars, histograms, and AI-guided highlights so model outputs become explorable visual evidence.",
        pdf_styles["Compact"],
    ),
    Paragraph("Brief Literature Survey", pdf_styles["Heading"]),
])

lit_table = Table(
    [[Paragraph("<b>Paper</b>", pdf_styles["SmallTable"]), Paragraph("<b>Main idea</b>", pdf_styles["SmallTable"]), Paragraph("<b>Use and limitation for ClariFi</b>", pdf_styles["SmallTable"])]]
    + [[Paragraph(c, pdf_styles["SmallTable"]) for c in row] for row in lit],
    colWidths=[1.15 * inch, 2.0 * inch, 3.1 * inch],
    repeatRows=1,
)
lit_table.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ("TOPPADDING", (0, 0), (-1, -1), 3),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
]))
story.append(lit_table)
story.append(Spacer(1, 4))

story.append(Paragraph("Plan of Activities", pdf_styles["Heading"]))
plan_table_pdf = Table(
    [[Paragraph(c, pdf_styles["SmallTable"]) for c in row] for row in plan],
    colWidths=[2.7 * inch, 1.55 * inch, 1.95 * inch],
    repeatRows=1,
)
plan_table_pdf.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F4F6F9")),
    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ("TOPPADDING", (0, 0), (-1, -1), 3),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
]))
story.append(plan_table_pdf)

story.append(PageBreak())
story.append(Paragraph("References", pdf_styles["Heading"]))
for ref in refs:
    story.append(Paragraph(ref, pdf_styles["Ref"]))

pdf_doc = SimpleDocTemplate(
    str(PDF_OUT),
    pagesize=letter,
    rightMargin=inch,
    leftMargin=inch,
    topMargin=inch,
    bottomMargin=inch,
)
pdf_doc.build(story)
print(PDF_OUT)
