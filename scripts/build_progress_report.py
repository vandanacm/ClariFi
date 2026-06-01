from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "Team07_ClariFi_Progress_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D7DE", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_run(run, bold=False, color=None, size=None):
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, bold=True, color="2E74B5", size=13)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    style_run(r, bold=True, color="1F4D78", size=11)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, bold=True, size=10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    return p


# ── Document setup ──────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.05

# ── Title block ─────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ClariFi — Project Progress Report")
style_run(r, bold=True, color="0B2545", size=17)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Team 07  |  Lalitha Dasu, Vandana Mansur, Pranav Manimaran  |  ECS 273 Visual Analytics  |  Spring 2026"
)
style_run(r, color="4B5563", size=10)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run("Date: 2026-05-24   |   Branch: main   |   Commits: 5")
style_run(r, color="6B7280", size=9.5)

doc.add_paragraph()  # spacer

# ── 1. What ClariFi Is ─────────────────────────────────────────────────────
add_section_heading(doc, "1. What ClariFi Is")
add_body(
    doc,
    "ClariFi is a personal finance and mortgage-readiness visual analytics application targeting "
    "California homebuyers. Users enter their income, debts, savings, and target property price; "
    "the app scores their readiness using a calibrated XGBoost ML model, compares them against "
    "real HMDA loan records, and an LLM agent (Ollama / Anthropic) explains the trade-offs in "
    "plain language through an interactive, multi-view dashboard.",
)

# ── 2. Architecture Overview ────────────────────────────────────────────────
add_section_heading(doc, "2. Architecture Overview")

arch_rows = [
    ["Layer", "Stack"],
    ["Frontend", "React 19, TypeScript, D3.js (SVG charts), Vite 6"],
    ["Backend", "FastAPI (Python 3), Uvicorn, port 8001"],
    ["ML Model", "XGBoost + isotonic calibration (scikit-learn Pipeline)"],
    ["LLM Agent", "Ollama (mistral, local) → Anthropic SDK fallback → rule-based fallback"],
    ["Data store", "Local JSON (public/data/local_store.json) + optional MongoDB"],
    ["Deployment", "Local dev (npm run dev:full)"],
]

arch_table = doc.add_table(rows=1, cols=2)
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(arch_table)
for i, text in enumerate(arch_rows[0]):
    arch_table.rows[0].cells[i].text = text
    set_cell_shading(arch_table.rows[0].cells[i], "E8F0FA")
    set_cell_margins(arch_table.rows[0].cells[i])
    for paragraph in arch_table.rows[0].cells[i].paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9.5)

for row in arch_rows[1:]:
    cells = arch_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(9.5)

arch_table.columns[0].width = Inches(1.4)
arch_table.columns[1].width = Inches(4.6)

# ── 3. Completed Features ───────────────────────────────────────────────────
add_section_heading(doc, "3. Completed Features")

# 3a Frontend
add_sub_heading(doc, "3a. Frontend")
frontend_items = [
    ("Landing page — ", "hero copy, feature cards, and call-to-action"),
    ("Onboarding flow — ", "collects income, debts, savings, target price, and market"),
    ("Authenticated dashboard — ", "8 interactive charts and a scenario control panel"),
    ("Dark/light theme toggle — ", "persisted via localStorage"),
    ("CSV transaction upload — ", "categorizes spending and back-fills scenario fields"),
    ("Auth modal — ", "register / login with JWT bearer token"),
]
for bold, rest in frontend_items:
    add_bullet(doc, rest, bold_prefix=bold)

# 3b Charts
add_sub_heading(doc, "3b. Charts (src/charts.tsx)")
chart_rows = [
    ["Chart", "Description", "Interactivity"],
    ["CashflowChart", "Waterfall bar: income → expenses → surplus", "Bars pop on hover; siblings dim; tooltip card"],
    ["IncomeHistogram", "Income distribution vs. HMDA cohort", "Crosshair line + income pill follows mouse"],
    ["ChoroplethMap", "California county readiness choropleth", "Smooth gradient legend; county tooltip (name, score, approval rate)"],
    ["RiskSurface", "DTI × down payment approval heatmap", '"You" marker with pulsing glow ring'],
    ["BenchmarkBars", "Income vs. BLS occupation benchmarks", "Static"],
    ["ExpenseDonut", "Spending breakdown donut", "Static"],
    ["HmdaScatter", "Loan amount vs. income scatter (HMDA)", "Static"],
    ["CalibrationChart", "Model predicted vs. actual approval rate", "Static"],
]

chart_table = doc.add_table(rows=1, cols=3)
chart_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(chart_table)
for i, text in enumerate(chart_rows[0]):
    chart_table.rows[0].cells[i].text = text
    set_cell_shading(chart_table.rows[0].cells[i], "E8F0FA")
    set_cell_margins(chart_table.rows[0].cells[i])
    for paragraph in chart_table.rows[0].cells[i].paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(8.8)

for row in chart_rows[1:]:
    cells = chart_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(8.5)

chart_table.columns[0].width = Inches(1.5)
chart_table.columns[1].width = Inches(2.3)
chart_table.columns[2].width = Inches(2.2)

# 3c Backend
add_sub_heading(doc, "3c. Backend API (api/main.py)")
api_rows = [
    ["Endpoint", "Purpose"],
    ["GET /api/health", "Liveness check; reports model + Ollama status"],
    ["POST /api/auth/register", "Create account (hashed password in local JSON store)"],
    ["POST /api/auth/login", "Returns JWT"],
    ["GET /api/me", "Decode current user from token"],
    ["GET/PUT /api/profile", "Persistent profile store"],
    ["POST /api/transactions/upload", "CSV ingestion → category totals"],
    ["GET /api/finance/summary", "Aggregated spending summary"],
    ["GET /api/benchmarks", "BLS occupation income benchmarks"],
    ["GET /api/hmda", "Processed HMDA loan data (California)"],
    ["GET /api/model", "Model report (AUC, calibration, feature importance)"],
    ["POST /api/mortgage/score", "Score a scenario with XGBoost; returns SHAP-style drivers"],
    ["POST/GET /api/scenarios", "Save / list user scenarios"],
    ["POST /api/agent/explain", "LLM explanation of readiness trade-offs"],
]

api_table = doc.add_table(rows=1, cols=2)
api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(api_table)
for i, text in enumerate(api_rows[0]):
    api_table.rows[0].cells[i].text = text
    set_cell_shading(api_table.rows[0].cells[i], "E8F0FA")
    set_cell_margins(api_table.rows[0].cells[i])
    for paragraph in api_table.rows[0].cells[i].paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for row in api_rows[1:]:
    cells = api_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(9)

api_table.columns[0].width = Inches(2.2)
api_table.columns[1].width = Inches(3.8)

# 3d ML Model
add_sub_heading(doc, "3d. ML Model")
ml_items = [
    ("Type: ", "XGBoost (300 estimators, depth 5) wrapped in isotonic CalibratedClassifierCV"),
    ("Training data: ", "58,000 synthetic HMDA-shaped California rows (4 counties)"),
    ("Features: ", "21 numeric (DTI, LTV, log-income, county relative metrics, flags) + 2 categorical (county code, loan type)"),
    ("Test AUC: ", "0.7191 (calibrated) / 0.8055 (raw XGBoost)"),
    ("Brier score: ", "0.1844 (calibrated)"),
    ("Artifact: ", "public/data/model_outputs/hmda_2025_xgboost_calibrated_pipeline.joblib (gitignored)"),
]
for bold, rest in ml_items:
    add_bullet(doc, rest, bold_prefix=bold)

# 3e LLM Agent
add_sub_heading(doc, "3e. LLM Agent")
agent_items = [
    ("Primary: ", "Ollama mistral at http://localhost:11434 — returns agentMode: \"mistral\""),
    ("Fallback 1: ", "Anthropic SDK (if ANTHROPIC_API_KEY set in .env)"),
    ("Fallback 2: ", "Rule-based template explanation"),
    ("Prompt: ", "Includes user scenario, score, SHAP drivers, and a free-form question"),
]
for bold, rest in agent_items:
    add_bullet(doc, rest, bold_prefix=bold)

# ── 4. Bugs Fixed ───────────────────────────────────────────────────────────
add_section_heading(doc, "4. Bugs Fixed")

bug_rows = [
    ["Bug", "Root Cause", "Fix"],
    [
        '"Static fallback" badge — model not scoring',
        "libomp.dylib (OpenMP) missing on macOS; XGBoost import failed silently",
        "brew install libomp; model artifact created by running training script",
    ],
    [
        "Model artifact missing",
        "scripts/train_xgboost_model.py had never been run",
        "Ran the script; saved .joblib to public/data/model_outputs/",
    ],
    [
        "Vite proxy wrong port",
        "vite.config.ts pointed to port 8080; API ran on 8001",
        "Updated proxy target to http://127.0.0.1:8001 in vite.config.ts and package.json",
    ],
    [
        "local_store.json corrupted",
        "Extra trailing byte caused JSONDecodeError: Extra data on every load_store() call",
        "Extracted valid JSON with JSONDecoder.raw_decode() and rewrote file",
    ],
    [
        "CSS animation on SVG r attribute",
        "@keyframes cannot animate SVG presentation attributes like r",
        "Switched pulse animation to transform: scale() with transform-box: fill-box",
    ],
    [
        "Ollama returning 500",
        "Caused by corrupted store file failing in the same load_store() call path",
        "Fixed after rewriting store",
    ],
]

bug_table = doc.add_table(rows=1, cols=3)
bug_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(bug_table)
for i, text in enumerate(bug_rows[0]):
    bug_table.rows[0].cells[i].text = text
    set_cell_shading(bug_table.rows[0].cells[i], "E8F0FA")
    set_cell_margins(bug_table.rows[0].cells[i])
    for paragraph in bug_table.rows[0].cells[i].paragraphs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(9)

for row in bug_rows[1:]:
    cells = bug_table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cells[i].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.size = Pt(8.8)

bug_table.columns[0].width = Inches(1.6)
bug_table.columns[1].width = Inches(2.3)
bug_table.columns[2].width = Inches(2.1)

# ── 5. Current System Status ────────────────────────────────────────────────
add_section_heading(doc, "5. Current System Status  (last verified 2026-05-24)")
status_items = [
    ("API: ", "online at http://127.0.0.1:8001"),
    ("Model: ", "calibrated-xgboost (confirmed via /api/mortgage/score)"),
    ("Ollama agent: ", "mistral (confirmed via /api/agent/explain)"),
    ("Frontend: ", "online at http://127.0.0.1:5173"),
    ("MongoDB: ", "not configured — optional; local JSON store is active"),
]
for bold, rest in status_items:
    add_bullet(doc, rest, bold_prefix=bold)

# ── 6. File Map ─────────────────────────────────────────────────────────────
add_section_heading(doc, "6. File Map")
file_map = doc.add_paragraph()
file_map.paragraph_format.space_after = Pt(4)
file_map.paragraph_format.line_spacing = 1.1
r = file_map.add_run(
    "ClariFi/\n"
    "├── api/\n"
    "│   ├── main.py              FastAPI app, 813 lines — all endpoints + ML inference\n"
    "│   └── requirements.txt     Python deps\n"
    "├── public/data/\n"
    "│   ├── local_store.json     Auth + scenarios + profiles (JSON file-based DB)\n"
    "│   ├── hmda_processed.json  Processed HMDA California loan data\n"
    "│   ├── bls_benchmarks.json  BLS occupation income benchmarks\n"
    "│   ├── model_report.json    Model metrics, calibration, feature importance\n"
    "│   └── model_outputs/       Gitignored — contains trained .joblib model\n"
    "├── scripts/\n"
    "│   ├── train_xgboost_model.py    Generate synthetic data + train XGBoost\n"
    "│   └── process_hmda_sample.mjs  Preprocess raw HMDA CSV\n"
    "├── src/\n"
    "│   ├── ReactApp.tsx         Main app shell, 997 lines\n"
    "│   ├── charts.tsx           All 8 D3 SVG chart components, 706 lines\n"
    "│   ├── styles.css           Global styles + theme vars + animations\n"
    "│   ├── types.ts             TypeScript interfaces\n"
    "│   ├── api.ts               Typed fetch client\n"
    "│   ├── Landing.tsx          Marketing landing page\n"
    "│   ├── Login.tsx            Auth modal\n"
    "│   └── Onboarding.tsx       Onboarding wizard\n"
    "├── vite.config.ts           Vite + React plugin + /api proxy → port 8001\n"
    "└── package.json             npm scripts: dev, dev:api, dev:full, build"
)
r.font.name = "Courier New"
r.font.size = Pt(8.5)

# ── 7. Immediate Next Steps ──────────────────────────────────────────────────
add_section_heading(doc, "7. Immediate Next Steps  (not yet built)")
next_items = [
    ("Scenario comparison — ", "save multiple scenarios and view them side-by-side"),
    ("Savings goal calculator — ", "\"how long until I can afford X% down on $Y price\""),
    ("Rate sensitivity — ", "show how monthly payment changes across 5.5–8% rates"),
    ("SHAP waterfall chart — ", "visualize per-feature model contributions for the current scenario"),
    ("MongoDB integration — ", "swap JSON file store for persistent cloud DB"),
    ("Deploy — ", "containerize API + serve Vite build (e.g., Fly.io + Vercel)"),
]
for bold, rest in next_items:
    add_bullet(doc, rest, bold_prefix=bold)

# ── Save ────────────────────────────────────────────────────────────────────
doc.core_properties.title = "ClariFi Progress Report"
doc.core_properties.author = "Team 07"
doc.core_properties.subject = "ECS 273 Visual Analytics — Spring 2026 Progress Report"
doc.save(OUT)
print(f"Saved: {OUT}")
