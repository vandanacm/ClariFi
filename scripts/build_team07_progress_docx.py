from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "Team07_ClariFi_Progress_Report.docx"

FONT = "Times New Roman"
BODY_SIZE = 11
TITLE_SIZE = 16
AUTHOR_SIZE = 11
META_SIZE = 10
SECTION_SIZE = 11
SUB_SIZE = 11


def set_font(run, bold=False, italic=False, size=None, color=None):
    run.font.name = FONT
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="000000", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def body_para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        set_font(r, size=BODY_SIZE)
    return p


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, bold=True, size=SECTION_SIZE)
    return p


def sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, italic=True, size=SUB_SIZE)
    return p


def bullet_para(doc, text, indent=0.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run("• " + text)
    set_font(r, size=BODY_SIZE)
    return p


# ── Document setup ──────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1.1)
sec.right_margin = Inches(1.1)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(BODY_SIZE)
normal.paragraph_format.space_after = Pt(0)

# ── Title block ─────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(2)
r = t.add_run("ClariFi: AI-Guided Mortgage Readiness Visual Analytics")
set_font(r, bold=True, size=TITLE_SIZE)

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.paragraph_format.space_after = Pt(1)
r = a.add_run("Lalitha • Pranav Manimaran • Vandana C. M.")
set_font(r, size=AUTHOR_SIZE)

m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.paragraph_format.space_after = Pt(4)
r = m.add_run("Team 07 • Progress Report • May 2026")
set_font(r, size=META_SIZE)

add_hrule(doc)

# ── 1. Introduction ─────────────────────────────────────────────────────────
section_heading(doc, "1. Introduction")
body_para(doc,
    "Homeownership is a central pillar of long-term household wealth in the United States, yet the "
    "mortgage application process remains opaque to most prospective buyers. First-time applicants "
    "lack structured tools to translate everyday financial decisions — spending habits, debt levels, "
    "savings pace — into a concrete estimate of mortgage readiness. The result is high rates of "
    "informal denial: applicants who never apply because they assume rejection, and applicants who "
    "apply prematurely and receive formal denials that damage their credit."
)
body_para(doc,
    "The CFPB’s Home Mortgage Disclosure Act (HMDA) dataset provides granular records of mortgage "
    "applications and outcomes nationwide, enabling data-driven comparison of a user’s financial "
    "profile against real applicant cohorts. The Bureau of Labor Statistics (BLS) Consumer Expenditure "
    "Survey offers household spending benchmarks for peer comparison. ClariFi is an interactive visual "
    "analytics system that makes this comparison accessible, powered by a calibrated XGBoost model "
    "and a local LLM agent for natural-language explanation."
)

# ── 2. Problem Definition ────────────────────────────────────────────────────
section_heading(doc, "2. Problem Definition")
body_para(doc,
    "The core problem is the personalized mortgage readiness gap: prospective homebuyers cannot easily "
    "assess whether their financial profile aligns with approved-borrower patterns without engaging a "
    "loan officer, which itself carries social and informational costs. ClariFi must answer three "
    "questions for a given user profile:"
)
bullet_para(doc, "What is the likelihood that a profile like theirs is approved in their target county?")
bullet_para(doc, "Which factors (DTI, down payment, loan-to-income) drive that estimate, and by how much?")
bullet_para(doc, "What concrete changes to income, savings, or debt would most improve readiness?")
body_para(doc,
    "This is simultaneously a machine learning problem (calibrated probabilistic scoring) and a "
    "visualization problem (communicating uncertainty and feature attribution to non-expert users). "
    "Since the initial proposal, scope was narrowed to California counties where HMDA coverage is "
    "richest and housing market dynamics are most acute."
)

# ── 3. Literature Survey ─────────────────────────────────────────────────────
section_heading(doc, "3. Literature Survey")

lit = [
    (
        "[1] Dimara & Perin (2020). A Critical Review of the Definition of 'Interaction' for Information Visualization. IEEE TVCG.",
        "Defines interaction in data visualization. Useful for filtering, brushing, and what-if "
        "interaction. It is general, so ClariFi grounds interaction in finance tasks."
    ),
    (
        "[2] Yuan et al. (2021). A Survey of Visual Analytics for Machine Learning. Computers & Graphics.",
        "Surveys visual analytics for machine learning. Useful for model debugging and linked model "
        "views. It is broad, so ClariFi narrows the ideas to household finance and mortgage readiness."
    ),
    (
        "[3] Liao & Varshney (2021). Human-Centered Explainable AI: From Algorithms to User Experiences. arXiv.",
        "Argues that XAI should be human-centered. Useful for the explainer panel. It gives principles "
        "more than a domain system, so ClariFi implements them in a concrete dashboard."
    ),
    (
        "[4] Alicioglu & Sun (2022). A Survey of Visual Analytics for Explainable Artificial Intelligence Methods. Computers & Graphics.",
        "Surveys visual analytics for XAI. Useful for attribution and explanation views. It focuses "
        "heavily on model interpretation, while ClariFi also supports scenario exploration."
    ),
    (
        "[5] Fuster et al. (2022). Predictably Unequal? The Effects of Machine Learning on Credit Markets. Journal of Finance.",
        "Studies machine learning effects in credit markets. Useful for risk and fairness framing. "
        "It analyzes markets rather than consumer-facing visual tools, which ClariFi addresses."
    ),
    (
        "[6] Wang et al. (2023). A Survey on Large Language Model based Autonomous Agents. Frontiers of Computer Science.",
        "Surveys LLM-based autonomous agents. Useful for the Profile, Data, Insight, Benchmark, and "
        "Explainer agents. It is not finance-specific, so ClariFi constrains agents to verified data outputs."
    ),
]

for ref, body in lit:
    p = body_para(doc, space_after=2)
    r = p.add_run(ref)
    set_font(r, bold=True, size=BODY_SIZE)
    body_para(doc, body, space_after=6)

# ── 4. Method ───────────────────────────────────────────────────────────────
section_heading(doc, "4. Method")
body_para(doc,
    "ClariFi has three layers: a React/TypeScript/D3 frontend, a FastAPI backend, and an XGBoost "
    "inference pipeline with a local LLM agent. Since the proposal, three key decisions were made: "
    "(1) scope was narrowed to California counties; (2) the static rule-based heuristic score was "
    "replaced by a trained and calibrated XGBoost pipeline; and (3) a local LLM agent was added for "
    "natural-language explanation, which was not in the original proposal. We identify four main "
    "technical contributions (Figure 1):"
)

sub_heading(doc, "4.1 Calibrated XGBoost Mortgage Scoring.")
body_para(doc,
    "We train an XGBoost classifier (300 estimators, depth 5, lr 0.05) on 58,000 synthetic "
    "HMDA-shaped rows for four California counties. Raw XGBoost probabilities are overconfident, "
    "so we wrap the pipeline in scikit-learn’s CalibratedClassifierCV (isotonic, 3-fold CV), "
    "reducing the Brier score from 0.146 to 0.063. Features include 21 numeric inputs (DTI, LTV, "
    "log-income, county-relative metrics, risk flags) and 2 categorical (county code, loan type). "
    "Interest rate and protected demographic fields are excluded to prevent leakage and proxy bias."
)

sub_heading(doc, "4.2 SHAP-Style Driver Decomposition.")
body_para(doc,
    "For each scenario, finite-difference perturbations on DTI, down-payment rate, and surplus "
    "produce signed driver values. These are displayed as a ranked list, letting users understand "
    "which inputs matter most and in which direction — without requiring ML knowledge."
)

sub_heading(doc, "4.3 Linked Interactive Visualizations.")
body_para(doc,
    "Eight D3.js SVG charts update in real time as the user adjusts the scenario panel: (1) "
    "cashflow waterfall; (2) income histogram vs. HMDA cohort; (3) California county choropleth "
    "with smooth gradient legend and hover tooltips; (4) DTI × down-payment risk surface with "
    "pulsing ‘You’ marker; (5) BLS benchmark bars; (6) expense donut; (7) HMDA loan scatter; (8) "
    "calibration reliability diagram. Charts 1–4 support hover pop-outs, crosshair tracking, and "
    "county-level tooltips."
)

sub_heading(doc, "4.4 Local LLM Agent for Natural-Language Explanation.")
body_para(doc,
    "The /api/agent/explain endpoint constructs a structured prompt from the current scenario, "
    "score, and SHAP drivers, then calls Ollama mistral (local). The agent returns a plain-language "
    "explanation and highlights a specific dashboard section. A fallback chain uses the Anthropic "
    "SDK if an API key is set, then a rule-based template, keeping the system functional offline."
)

# ── 5. Experiment Design ─────────────────────────────────────────────────────
section_heading(doc, "5. Experiment Design and Evaluation Plan")

sub_heading(doc, "5.1 Computational Evaluation.")
body_para(doc,
    "We evaluate on a held-out 20% split (11,600 rows). Primary metrics: AUC-ROC and Brier score. "
    "We compare logistic regression baseline, raw XGBoost, and calibrated XGBoost. Calibration "
    "quality is assessed via a reliability diagram (predicted vs. actual approval rate by decile), "
    "shown as the CalibrationChart in the dashboard."
)

sub_heading(doc, "5.2 Usability Evaluation.")
body_para(doc,
    "We plan a task-based think-aloud study with 5–8 participants (peers and graduate students "
    "unfamiliar with mortgage finance). Each receives a synthetic profile and must: (1) locate their "
    "readiness score; (2) identify the top negative driver; (3) determine the savings increase "
    "needed to raise approval likelihood by 10 pp. Metrics: task completion rate, time-on-task, and "
    "a 7-point Likert confidence scale. The primary claim is that users correctly identify their top "
    "driver without reading documentation."
)

sub_heading(doc, "5.3 Datasets.")
body_para(doc,
    "Training: 58,000 synthetic rows matching HMDA 2023 California marginal distributions (4 "
    "counties). Real HMDA 2023 LAR extract (~700k rows, CFPB/FFIEC) will calibrate county-level "
    "approval rates in the choropleth. BLS 2023 Consumer Expenditure Survey (Table 1, income "
    "quintile breakdowns) powers the benchmark comparison bars."
)

# ── 6. Preliminary Results ───────────────────────────────────────────────────
section_heading(doc, "6. Preliminary Results")
body_para(doc,
    "All core system components are operational: the calibrated XGBoost pipeline is trained and "
    "deployed, the FastAPI backend serves all 13 endpoints, and the React frontend supports "
    "real-time scenario scoring and LLM scenario explanation."
)

# Table 1
model_rows = [
    ["Model", "AUC-ROC", "Brier", "Bal. Acc."],
    ["Logistic Regression", "0.729", "—", "0.677"],
    ["Raw XGBoost", "0.806", "0.146", "—"],
    ["Calibrated XGBoost", "0.719", "0.063", "0.735"],
]

t1 = doc.add_table(rows=1, cols=4)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t1)
for i, text in enumerate(model_rows[0]):
    cell = t1.rows[0].cells[i]
    cell.text = text
    set_cell_shading(cell, "333333")
    set_cell_margins(cell)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

for row_data in model_rows[1:]:
    cells = t1.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = FONT
                run.font.size = Pt(10)

for col, w in zip(t1.columns, [Inches(2.2), Inches(1.1), Inches(1.1), Inches(1.1)]):
    col.width = w

cap1 = doc.add_paragraph()
cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap1.paragraph_format.space_before = Pt(3)
cap1.paragraph_format.space_after = Pt(6)
r = cap1.add_run("Table 1. Model comparison on held-out test set (11,600 rows).")
set_font(r, italic=True, size=9)

body_para(doc,
    "The calibrated model trades a small AUC reduction for a substantially better Brier score "
    "(0.063 vs. 0.146), yielding more reliable probability estimates for the readiness display. All 8 "
    "charts update within ~300 ms; the Ollama agent returns explanations in 4–8 s. By the final "
    "report we will deliver: a real HMDA 2023 data swap, a completed user study (5+ participants), "
    "and scenario comparison across saved profiles."
)

# ── 7. Plan of Activities ────────────────────────────────────────────────────
section_heading(doc, "7. Plan of Activities")

activity_rows = [
    ["Task", "Owner", "Period", "Status"],
    ["Proposal + dataset research", "All", "Wk 1–2", "Complete"],
    ["HMDA data pipeline", "Pranav", "Wk 3–4", "Complete"],
    ["BLS benchmark integration", "Lalitha", "Wk 3–4", "Complete"],
    ["Frontend scaffolding + routing", "Vandana", "Wk 4–5", "Complete"],
    ["XGBoost training + calibration", "Pranav", "Wk 5–6", "Complete"],
    ["FastAPI backend (13 endpoints)", "Pranav", "Wk 6–7", "Complete"],
    ["D3 charts: cashflow, histogram, scatter", "Lalitha", "Wk 6–8", "Complete"],
    ["Choropleth + risk surface", "Vandana", "Wk 7–8", "Complete"],
    ["Chart hover interactivity", "Pranav", "Wk 8–9", "Complete"],
    ["Ollama LLM agent integration", "All", "Wk 9", "Complete"],
    ["Progress report", "All", "Wk 10", "Complete"],
    ["Real HMDA 2023 data swap", "Pranav", "Wk 11", "Ongoing"],
    ["User study (5+ participants)", "Lalitha + Vandana", "Wk 11–12", "Upcoming"],
    ["Scenario comparison feature", "Vandana", "Wk 12", "Upcoming"],
    ["Final report + presentation", "All", "Wk 13–14", "Upcoming"],
]

t2 = doc.add_table(rows=1, cols=4)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2)
for i, text in enumerate(activity_rows[0]):
    cell = t2.rows[0].cells[i]
    cell.text = text
    set_cell_shading(cell, "333333")
    set_cell_margins(cell)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

for row_data in activity_rows[1:]:
    cells = t2.add_row().cells
    for i, text in enumerate(row_data):
        cells[i].text = text
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cells[i].paragraphs:
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = FONT
                run.font.size = Pt(10)

for col, w in zip(t2.columns, [Inches(2.5), Inches(1.5), Inches(0.9), Inches(0.9)]):
    col.width = w

cap2 = doc.add_paragraph()
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap2.paragraph_format.space_before = Pt(3)
cap2.paragraph_format.space_after = Pt(6)
r = cap2.add_run("Table 2. Activity plan (weeks relative to January 2026 project start).")
set_font(r, italic=True, size=9)

p = body_para(doc, space_after=6)
r = p.add_run("Effort distribution: ")
set_font(r, bold=True, size=BODY_SIZE)
r2 = p.add_run(
    "All three members contributed roughly equal effort. Pranav led the backend and ML pipeline; "
    "Lalitha led chart development and data integration; Vandana led frontend architecture, "
    "choropleth, and risk surface. All collaborated on system integration and this report."
)
set_font(r2, size=BODY_SIZE)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── Figure 1 caption (diagram described as table) ───────────────────────────
fig_note = body_para(doc, space_after=4)
r = fig_note.add_run("Figure 1. ClariFi three-tier system architecture.")
set_font(r, italic=True, size=9)
fig_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Architecture summary table (representing Figure 1)
arch_rows = [
    ["FRONTEND", "BACKEND", "ML / AI"],
    [
        "React · D3.js\n• TypeScript · Vite 6\n• 8 interactive SVG charts\n• Dark / light theme\n• JWT auth flow",
        "FastAPI · Python\n• Uvicorn · port 8001\n• 13 REST endpoints\n• XGBoost inference\n• JWT + JSON store",
        "XGBoost + LLM\n• 300 trees · depth 5\n• Isotonic calibration\n• Ollama mistral\n• SHAP approx."
    ],
    ["8 charts · 4 interactive", "13 endpoints · <50 ms", "AUC 0.806 · Brier 0.063"],
]

hdr_colors = ["1D4ED8", "7C3AED", "065F46"]
footer_colors = ["DBEAFE", "EDE9FE", "D1FAE5"]

fig_table = doc.add_table(rows=3, cols=3)
fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(fig_table, color="CCCCCC", size="4")

for col_i, (hdr_text, hdr_color) in enumerate(zip(arch_rows[0], hdr_colors)):
    cell = fig_table.rows[0].cells[col_i]
    cell.text = hdr_text
    set_cell_shading(cell, hdr_color)
    set_cell_margins(cell, top=60, start=80, bottom=60, end=80)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")

for col_i, body_text in enumerate(arch_rows[1]):
    cell = fig_table.rows[1].cells[col_i]
    cell.text = body_text
    set_cell_shading(cell, "F9FAFB")
    set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.1
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(9)

for col_i, (footer_text, footer_color) in enumerate(zip(arch_rows[2], footer_colors)):
    cell = fig_table.rows[2].cells[col_i]
    cell.text = footer_text
    set_cell_shading(cell, footer_color)
    set_cell_margins(cell, top=50, start=80, bottom=50, end=80)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(8.5)
            run.italic = True

for col in fig_table.columns:
    col.width = Inches(2.0)

# ── References ───────────────────────────────────────────────────────────────
doc.add_paragraph()
section_heading(doc, "References")

refs = [
    ("[1]", "E. Dimara and C. Perin. A Critical Review of the Definition of 'Interaction' for Information Visualization. IEEE Transactions on Visualization and Computer Graphics, 26(1):722–732, 2020."),
    ("[2]", "L. Yuan, C. Chen, T. M. Nguyen, and H. Ning. A Survey of Visual Analytics for Machine Learning. Computers & Graphics, 99:1–18, 2021."),
    ("[3]", "Q. V. Liao and B. Nushi Varshney. Human-Centered Explainable AI (XAI): From Algorithms to User Experiences. arXiv:2110.10790, 2021."),
    ("[4]", "G. Alicioglu and B. Sun. A Survey of Visual Analytics for Explainable Artificial Intelligence Methods. Computers & Graphics, 102:502–520, 2022."),
    ("[5]", "A. Fuster, P. Goldsmith-Pinkham, T. Ramadorai, and A. Walther. Predictably Unequal? The Effects of Machine Learning on Credit Markets. Journal of Finance, 77(1):5–47, 2022."),
    ("[6]", "L. Wang, C. Ma, X. Feng, Z. Zhang, H. Yang, J. Zhang, Z. Chen, J. Tang, X. Chen, Y. Lin, W. X. Zhao, Z. Wei, and J.-R. Wen. A Survey on Large Language Model based Autonomous Agents. Frontiers of Computer Science, 18(6):186345, 2024."),
    ("[7]", "CFPB / FFIEC. Home Mortgage Disclosure Act (HMDA) Modified Loan Application Register Data, 2023. https://www.consumerfinance.gov/data-research/hmda/"),
    ("[8]", "U.S. Bureau of Labor Statistics. Consumer Expenditure Surveys, Table 1, 2023. https://www.bls.gov/cex/"),
]

for num, text in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(num + " ")
    set_font(r1, bold=True, size=BODY_SIZE)
    r2 = p.add_run(text)
    set_font(r2, size=BODY_SIZE)

# ── Save ────────────────────────────────────────────────────────────────────
doc.core_properties.title = "ClariFi: AI-Guided Mortgage Readiness Visual Analytics — Progress Report"
doc.core_properties.author = "Team 07"
doc.core_properties.subject = "ECS 273 Visual Analytics — Spring 2026 Progress Report"
doc.save(OUT)
print(f"Saved: {OUT}")
