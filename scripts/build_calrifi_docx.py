from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "CalriFi_proposal_report.docx"

ACCENT = RGBColor(0x0A, 0x69, 0x68)
MUTED = RGBColor(0x5D, 0x6B, 0x7A)
LIGHT = "EAF3F2"
BORDER = "C9D6D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text, bold=False, color=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_q(doc, label, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    add_run(p, label, bold=True, color=ACCENT)
    body = doc.add_paragraph(text)
    body.style = "Body Text"


def add_ref(doc, text):
    p = doc.add_paragraph()
    p.style = "Reference"
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)

styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(18)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(0x16, 0x24, 0x33)

styles["Subtitle"].font.name = "Arial"
styles["Subtitle"].font.size = Pt(10.5)
styles["Subtitle"].font.color.rgb = MUTED

styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(13)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = ACCENT

styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(11)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = ACCENT

body_style = styles["Body Text"] if "Body Text" in styles else styles.add_style("Body Text", 1)
body_style.font.name = "Arial"
body_style.font.size = Pt(11)
body_style.paragraph_format.space_after = Pt(5)
body_style.paragraph_format.line_spacing = 1.04

ref_style = styles["Reference"] if "Reference" in styles else styles.add_style("Reference", 1)
ref_style.font.name = "Arial"
ref_style.font.size = Pt(10)
ref_style.paragraph_format.space_after = Pt(4)
ref_style.paragraph_format.line_spacing = 1.0

header = section.header.paragraphs[0]
header.text = "CalriFi Proposal"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Arial"
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = MUTED

title = doc.add_paragraph()
title.style = "Title"
title.add_run("CalriFi: Interactive Visual Analytics for Mortgage Readiness")

subtitle = doc.add_paragraph()
subtitle.style = "Subtitle"
subtitle.add_run("Team 7 | Lalitha Dasu, Vandana Mansur, Pranav Manimaran")

intro = doc.add_paragraph()
intro.style = "Body Text"
intro.add_run(
    "This proposal describes a visual analytics project that combines personal finance inputs with real-world mortgage application data. "
    "The emphasis is on objective data exploration, model-based analysis, and interactive visualization, rather than personal financial advice."
)

doc.add_heading("Heilmeier Questions", level=1)

add_q(
    doc,
    "Q1. What are you trying to do?",
    "CalriFi helps individuals understand mortgage readiness by comparing a personal financial profile, including income, savings, and debt, against real loan application records from the Home Mortgage Disclosure Act (HMDA) dataset. Users will see where they stand relative to approved and denied borrowers in a target geography and can simulate feasible changes, such as a larger down payment or lower monthly debt, to see how those changes affect estimated approval likelihood.",
)

add_q(
    doc,
    "Q2. How is it done today, and what are the limits?",
    "Today, most consumer tools separate the problem into either budgeting or mortgage affordability. Budgeting tools track cash flow, while mortgage calculators mainly use formula-based estimates such as debt-to-income thresholds. HMDA tools expose public lending data, but they are not designed around a user's own financial profile. Academic work studies mortgage prediction and fair-lending questions, but usually does not provide an interactive interface for personal exploration. FinSight, our reference architecture, forecasts savings, but it does not connect a user's financial trajectory to real mortgage outcomes. The gap is a system that combines personal finance, large-scale mortgage data, model explanations, and linked visual exploration.",
)

add_q(
    doc,
    "Q3. What is new in your approach? Why will it succeed?",
    "CalriFi adds three main pieces. First, we will train an approval-likelihood model on a reasonably large HMDA subset, initially focused on 2018-2022 or the most recent annual California records depending on download size and preprocessing time. We plan to use logistic regression as a baseline and XGBoost as the main nonlinear model, with SHAP feature attributions for explanation. Second, we will build a what-if engine that changes only user-controllable fields, such as debt, income, savings, and target loan amount, to estimate which changes matter most. Third, we will build a coordinated visual analytics interface with a geographic view, borrower comparison plot, linked distributions, and scenario controls. This should succeed because the interface grounds personal finance decisions in real application patterns instead of a generic calculator.",
)

add_q(
    doc,
    "Q4. Who cares?",
    "Primary users are first-time homebuyers and renters who want to understand the gap between their current finances and common mortgage approval patterns. Secondary stakeholders include mortgage counselors, consumer fintech teams, and researchers interested in geographic differences in credit access. For the course, the project is relevant because it requires large real-world data, algorithmic analysis, and interaction with analysis results.",
)

doc.add_heading("Literature Survey", level=1)
lit_items = [
    "[1] Fuster et al. (2022) study how machine learning changes mortgage credit allocation using HMDA-like credit market data. This directly informs our modeling and fairness caution. Its limitation for us is that it is policy-focused and does not provide an interactive visual analytics tool.",
    "[2] Chen and Guestrin (2016) introduce XGBoost, a scalable tree boosting method. It is useful for modeling nonlinear relationships between income, DTI, loan amount, geography, and approval outcomes. Its limitation is that it does not provide native counterfactual explanations.",
    "[3] Wachter et al. (2017) define counterfactual explanations as actionable changes that could alter a model outcome. This supports our what-if engine. The limitation is computational cost in high dimensions, so we will restrict scenarios to user-controllable features.",
    "[4] Keim et al. (2008) define visual analytics as the combination of automated analysis and human interaction. This gives the conceptual basis for our linked analysis workflow, though it is not specific to mortgage or finance data.",
    "[5] Wongsuphasawat et al. (2016) present Voyager, which supports exploratory visual analysis through recommendations and faceted browsing. It inspires our coordinated exploration design, but it is general-purpose and does not address geospatial mortgage data.",
]
for item in lit_items:
    p = doc.add_paragraph(style="Body Text")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.add_run(item)

add_q(
    doc,
    "Q5. What difference will it make, and how do we measure success?",
    "Success means users can answer questions that are difficult to answer from a calculator alone: which regions are more realistic, whether savings or debt payoff matters more, and why a model score changes under a scenario. We will measure model quality using held-out HMDA labels with AUC-ROC, F1, calibration, and comparison to a logistic regression baseline. We will measure visualization utility with a small task-based study of about five users, tracking task completion, correctness, and confidence for questions such as identifying the single most useful change for approval likelihood.",
)

add_q(
    doc,
    "Q6. What are the risks and payoffs?",
    "Risks include HMDA labels reflecting historical lender behavior and possible bias, messy categorical fields, and users over-interpreting a model estimate as a real loan decision. We will mitigate these risks by documenting feature limitations, filtering to a clear mortgage application type, showing uncertainty and explanations, and avoiding claims of guaranteed approval. The payoff is a system that is more than a dashboard: it combines large-scale public data, predictive modeling, explainability, and interactive visual analysis.",
)

add_q(
    doc,
    "Q7. How much will it cost?",
    "The financial cost is zero. HMDA data are freely available from CFPB/FFIEC, FinSight is open-source, and our implementation will use open-source Python and JavaScript tools. The main cost is team time and local or free-tier cloud compute for preprocessing and model training.",
)

add_q(
    doc,
    "Q8. How long will it take?",
    "The project will take six weeks. We already have an early frontend prototype and a sample data-processing pipeline. The remaining work is to replace the sample with real HMDA data, train and evaluate the model, implement linked interactions, conduct a small user study, and prepare the final report and presentation.",
)

add_q(
    doc,
    "Q9. What are the midterm and final checks for success?",
    "By Week 3, we should have HMDA data cleaned, a baseline model trained with target AUC around 0.75 or higher, and a static county-level approval visualization. By Week 6, we should have a complete interactive dashboard with what-if simulation, profile integration, brushing-and-linking across views, model explanations, and user study results.",
)

doc.add_heading("Plan of Activities", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
headers = ["Activity", "Owner(s)", "Weeks"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = text
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(9.5)

rows = [
    ("HMDA data acquisition, cleaning, and EDA", "All", "1-2"),
    ("FinSight reference setup and CalriFi integration planning", "Vandana, Pranav", "1-2"),
    ("XGBoost model training and SHAP analysis", "Lalitha, Vandana", "2-3"),
    ("Counterfactual what-if engine", "Lalitha", "3-4"),
    ("Geographic approval-rate visualization", "Pranav", "3-4"),
    ("Brushing-and-linking dashboard interactions", "All", "4-5"),
    ("Profile-to-mortgage readiness integration", "Vandana", "5"),
    ("User study and analysis", "All", "5-6"),
    ("Final report and presentation", "All", "6"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        cells[i].text = text
        set_cell_border(cells[i])
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[i].paragraphs:
            p.style = "Body Text"
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
        if i == 2:
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_table_width(table, [5000, 2860, 1500])

doc.add_section(WD_SECTION_START.NEW_PAGE)
doc.add_heading("References", level=1)
references = [
    "[1] Fuster, A., Goldsmith-Pinkham, P., Ramadorai, T., and Walther, A. (2022). Predictably unequal? The effects of machine learning on credit markets. Journal of Finance, 77(1), 5-47.",
    "[2] Chen, T., and Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.",
    "[3] Wachter, S., Mittelstadt, B., and Russell, C. (2017). Counterfactual explanations without opening the black box. Harvard Journal of Law and Technology, 31(2), 841-887.",
    "[4] Keim, D., Andrienko, G., Fekete, J.-D., Goerg, C., Kohlhammer, J., and Melancon, G. (2008). Visual analytics: Definition, process, and challenges. In Information Visualization, LNCS 4950, 154-175.",
    "[5] Wongsuphasawat, K., Moritz, D., Anand, A., Mackinlay, J., Howe, B., and Heer, J. (2016). Voyager: Exploratory analysis via faceted browsing. IEEE Transactions on Visualization and Computer Graphics, 22(1), 649-658.",
    "[6] FinSight. (2024). A Smart Visual Analytics Platform for Personal Finance Tracking and Goal-Oriented Investment Planning. GitHub. https://github.com/haseebshaik00/FinSight",
    "[7] Consumer Financial Protection Bureau and Federal Financial Institutions Examination Council. HMDA Data Publication Platform. https://ffiec.cfpb.gov/",
]
for ref in references:
    add_ref(doc, ref)

doc.save(OUT)
print(OUT)
